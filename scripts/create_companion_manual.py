from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "doc"
TMP = ROOT / "tmp" / "docs" / "xiaolu_manual"
SHEET = ROOT / "assets" / "xiaolu" / "spritesheet.webp"
BOOKMARKS = ROOT / "assets" / "bookmarks"
DOCX_PATH = OUT / "共学日记说明书.docx"

PURPLE = "76558F"
PURPLE_DARK = "4B315E"
LILAC = "E9DDF1"
CREAM = "FFF8E9"
PAPER = "FFFDF7"
PEACH = "F4CFA7"
ROSE = "EFA6B4"
GREEN = "A9C99E"
YELLOW = "F4D57B"
INK = "3A2C3E"
MUTED = "756A77"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = PURPLE_DARK, size: int = 14, sides: Iterable[str] = ("top", "left", "bottom", "right")) -> None:
    # Word resolves adjacent cell borders independently. Keeping every ordinary
    # table edge at one pixel weight prevents thicker or broken-looking left
    # edges after PDF export; only the cover keeps its heavy frame.
    size = 24 if size >= 24 else 8
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in sides:
        tag = f"w:{side}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, centimeters: float) -> None:
    cell.width = Cm(centimeters)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(centimeters / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size: float, bold: bool = False, color: str = INK, name: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, size: float = 9, bold: bool = False, color: str = INK, name: str = "Microsoft YaHei"):
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, color, name)
    return run


def set_para(paragraph, before: float = 0, after: float = 0, line: float = 1.15, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def clear_cell(cell) -> None:
    cell.text = ""
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_label(doc: Document, english: str, chinese: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    left, right = table.rows[0].cells
    set_cell_width(left, 2.55)
    set_cell_width(right, 9.35)
    clear_cell(left)
    clear_cell(right)
    set_cell_shading(left, PURPLE_DARK)
    set_cell_shading(right, LILAC)
    set_cell_border(left, PURPLE_DARK, 16)
    set_cell_border(right, PURPLE_DARK, 16)
    p = left.paragraphs[0]
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, english, 6.5, True, WHITE, "Consolas")
    p = right.paragraphs[0]
    set_para(p)
    add_text(p, chinese, 13, True, PURPLE_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_footer(section, page_no: str) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p)
    add_text(p, "XIAOLU STUDY JOURNAL  ·  ", 6.5, True, PURPLE, "Consolas")
    run = p.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_1, instr_text, fld_char_2))
    set_run_font(run, 6.5, True, PURPLE, "Consolas")


def add_page(doc: Document, page_no: str) -> None:
    doc.add_page_break()
    add_footer(doc.sections[-1], page_no)


def extract_sprites() -> dict[str, Path]:
    TMP.mkdir(parents=True, exist_ok=True)
    source = Image.open(SHEET).convert("RGBA")
    cw, ch = source.width // 8, source.height // 11
    specs = {
        "idle": (0, 2),
        "look": (9, 0),
        "run_right": (1, 0),
        "run_left": (2, 0),
        "wave": (3, 2),
        "jump": (4, 2),
        "failed": (5, 5),
        "waiting": (6, 3),
        "focus": (7, 3),
        "review": (8, 4),
    }
    results: dict[str, Path] = {}
    for name, (row, col) in specs.items():
        frame = source.crop((col * cw, row * ch, (col + 1) * cw, (row + 1) * ch))
        alpha = frame.getchannel("A")
        box = alpha.getbbox()
        if box:
            frame = frame.crop(box)
        # A nearest-neighbour enlargement keeps the original pixel edges crisp.
        frame = frame.resize((frame.width * 3, frame.height * 3), Image.Resampling.NEAREST)
        path = TMP / f"sprite_{name}.png"
        frame.save(path)
        results[name] = path
    return results


def make_pixel_icon(label: str, fill: str, filename: str) -> Path:
    canvas = Image.new("RGBA", (360, 160), (255, 253, 247, 0))
    d = ImageDraw.Draw(canvas)
    d.rectangle((12, 12, 348, 148), fill="#4B315E")
    d.rectangle((20, 20, 340, 140), fill=f"#{fill}")
    # Decorative pixels; text stays in Word so Chinese rendering is reliable.
    for x, y in ((30, 30), (310, 30), (30, 110), (310, 110)):
        d.rectangle((x, y, x + 18, y + 18), fill="#FFF8E9")
    path = TMP / filename
    canvas.save(path)
    return path


def add_note_box(doc: Document, title: str, body: str, fill: str = CREAM, accent: str = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    clear_cell(cell)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, 14)
    p = cell.paragraphs[0]
    set_para(p, after=3)
    add_text(p, title, 9.5, True, accent)
    p = cell.add_paragraph()
    set_para(p, line=1.25)
    add_text(p, body, 8.3, False, INK)


def add_action_card(cell, image: Path, title: str, body: str, fill: str) -> None:
    clear_cell(cell)
    set_cell_shading(cell, fill)
    set_cell_border(cell, PURPLE_DARK, 10)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=1)
    p.add_run().add_picture(str(image), height=Cm(2.0))
    p = cell.add_paragraph()
    set_para(p, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, title, 8.4, True, PURPLE_DARK)
    p = cell.add_paragraph()
    set_para(p, line=1.12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, body, 6.8, False, MUTED)


def add_bullet(doc: Document, title: str, body: str, color: str = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    dot, text = table.rows[0].cells
    set_cell_width(dot, 0.72)
    set_cell_width(text, 11.08)
    clear_cell(dot)
    clear_cell(text)
    set_cell_shading(dot, color)
    set_cell_border(dot, PURPLE_DARK, 8)
    p = dot.paragraphs[0]
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "+", 9, True, WHITE, "Consolas")
    p = text.paragraphs[0]
    set_para(p, line=1.18)
    add_text(p, title + "  ", 8.4, True, INK)
    add_text(p, body, 7.35, False, MUTED)


def build_document() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    sprites = extract_sprites()
    make_pixel_icon("", PEACH, "pixel_panel.png")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(14.8)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    # Cover
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover.cell(0, 0)
    clear_cell(cell)
    set_cell_shading(cell, CREAM)
    set_cell_border(cell, PURPLE_DARK, 24)
    set_cell_margins(cell, 220, 180, 220, 180)
    p = cell.paragraphs[0]
    set_para(p, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "XIAOLU STUDY JOURNAL", 8, True, PURPLE, "Consolas")
    p = cell.add_paragraph()
    set_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "共学日记说明书", 23, True, PURPLE_DARK)
    p = cell.add_paragraph()
    set_para(p, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "一份只属于我们两个人的学习约定", 9.2, False, MUTED)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(sprites["wave"]), height=Cm(7.1))
    p = cell.add_paragraph()
    set_para(p, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "她不是宠物。", 11, True, PURPLE_DARK)
    p = cell.add_paragraph()
    set_para(p, line=1.35, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "她是住在桌面上的学习搭子，\n替现实里的你来陪我、提醒我，也见证我们一起认真过的每一天。", 9, False, INK)
    p = doc.add_paragraph()
    set_para(p, before=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "VERSION 0.3.0  ·  2026", 6.8, True, PURPLE, "Consolas")
    add_footer(section, "00")

    # Page 1: identity and controls
    add_page(doc, "01")
    add_label(doc, "HELLO", "先认识一下小鹿")
    intro = doc.add_table(rows=1, cols=2)
    intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    intro.autofit = False
    c0, c1 = intro.rows[0].cells
    set_cell_width(c0, 4.15); set_cell_width(c1, 7.65)
    clear_cell(c0); clear_cell(c1)
    set_cell_shading(c0, LILAC); set_cell_border(c0, PURPLE_DARK, 12)
    p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(sprites["idle"]), height=Cm(4.35))
    set_cell_shading(c1, PAPER); set_cell_border(c1, PURPLE_DARK, 12)
    p = c1.paragraphs[0]; set_para(p, after=4)
    add_text(p, "为什么她会在这里？", 10.5, True, PURPLE_DARK)
    p = c1.add_paragraph(); set_para(p, line=1.35)
    add_text(p, "因为现实里的我们约好要互相监督学习。电脑里的小鹿代表你：她会在关键时间问我有没有到位，也会把每天的努力收进一本共学日记。", 8.2)
    p = c1.add_paragraph(); set_para(p, before=4, line=1.25)
    add_text(p, "所有记录只保存在这台电脑里，不会自动发给别人。", 7.5, True, PURPLE)

    p = doc.add_paragraph(); set_para(p, before=7, after=4)
    add_text(p, "三个最常用的操作", 11, True, PURPLE_DARK)
    controls = doc.add_table(rows=3, cols=1)
    controls.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ("双击小鹿", "开始或结束一段学习计时。一天可以分成很多段，最后自动相加。", GREEN),
        ("右键小鹿", "打开像素风“共学日记”，看今日记录、历史统计与书签收藏。", PEACH),
        ("拖动小鹿", "把她挪到不挡视线的位置。她会朝拖动方向小跑，松手后恢复平常。", LILAC),
    ]
    for row, (title, body, fill) in zip(controls.rows, items):
        cell = row.cells[0]; clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 10)
        p = cell.paragraphs[0]; set_para(p, after=1)
        add_text(p, title + "  ", 9, True, PURPLE_DARK)
        add_text(p, body, 7.7, False, INK)
        no_split(row)
    add_note_box(doc, "平常的她", "没有计时时，小鹿会安静待着，也会顺着鼠标方向看过来。暂停不用登记；想学时再双击就好。", CREAM, PURPLE)

    # Page 2: schedule and check-in rules
    add_page(doc, "02")
    add_label(doc, "CHECK IN", "一天五次，只确认“我在”")
    p = doc.add_paragraph(); set_para(p, after=5, line=1.3)
    add_text(p, "在约定时间，小鹿会弹出一句提醒。点一下“我在”，就是告诉她：这个时间我在学习现场。它和学习计时互不干扰。", 8.4)

    schedule = doc.add_table(rows=6, cols=3)
    schedule.alignment = WD_TABLE_ALIGNMENT.CENTER
    schedule.autofit = False
    widths = [Cm(2.0), Cm(3.1), Cm(6.75)]
    for col, width in zip(schedule.columns, widths):
        col.width = width
    headers = ("时间", "有效窗口", "这一刻的意义")
    for i, text in enumerate(headers):
        cell = schedule.rows[0].cells[i]; clear_cell(cell); set_cell_shading(cell, PURPLE_DARK); set_cell_border(cell, PURPLE_DARK, 10)
        p = cell.paragraphs[0]; set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, text, 7.5, True, WHITE)
    set_repeat_table_header(schedule.rows[0])
    rows = [
        ("09:00", "08:55-09:05", "早上的开始打卡：先确认到位，真正计时仍由双击决定。"),
        ("12:00", "11:55-12:05", "中午在场确认。"),
        ("15:00", "14:55-15:05", "下午在场确认。"),
        ("18:00", "17:55-18:05", "傍晚在场确认。"),
        ("21:00", "20:55-21:05", "结束打卡：确认后打开今日结算。"),
    ]
    fills = [CREAM, PAPER, CREAM, PAPER, LILAC]
    for row, data, fill in zip(schedule.rows[1:], rows, fills):
        no_split(row)
        for idx, text in enumerate(data):
            cell = row.cells[idx]; clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 7)
            p = cell.paragraphs[0]; set_para(p, line=1.18, align=WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT)
            add_text(p, text, 7.3 if idx == 2 else 7.7, idx == 0, PURPLE_DARK if idx == 0 else INK, "Consolas" if idx < 2 else "Microsoft YaHei")

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_note_box(doc, "十分钟规则", "每次打卡只有正负五分钟的有效期。超时没有点，就会记为“未打卡”；不能普通补签，但也不会打断正在进行的学习计时。", "FCE7E8", ROSE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_note_box(doc, "开机就会在", "小鹿会随 Windows 登录自动启动。若电脑在有效窗口内刚开机或刚唤醒，仍可以正常打卡；窗口已经结束，则按未打卡记录。", "E8F1E5", GREEN)

    # Page 3: speech bubbles
    add_page(doc, "03")
    add_label(doc, "SAY HELLO", "提醒会像朋友说话，不像闹钟")
    p = doc.add_paragraph(); set_para(p, after=6, line=1.3)
    add_text(p, "同一个提醒会从多套台词里随机挑一句。语气亲近、简短，不催命，也不会每天一模一样。下面只是几种例子。", 8.3)

    bubbles = [
        ("09:00", "早呀，我来啦。你也到位了吗？", CREAM),
        ("09:00", "九点啦，一起把今天开个好头吧。", LILAC),
        ("12:00", "到中午啦，给我一个“我在”好不好？", PEACH),
        ("15:00", "我来偷偷看一眼，你还在认真吗？", "E8F1E5"),
        ("18:00", "六点报到！今天也坚持到这里啦。", CREAM),
        ("21:00", "今天辛苦啦，要不要和我一起收个尾？", LILAC),
    ]
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (time, quote, fill) in enumerate(bubbles):
        cell = table.rows[idx // 2].cells[idx % 2]
        clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 10)
        p = cell.paragraphs[0]; set_para(p, after=3)
        add_text(p, time, 7.2, True, PURPLE, "Consolas")
        p = cell.add_paragraph(); set_para(p, line=1.3)
        add_text(p, "“" + quote + "”", 8.2, True, INK)
        no_split(table.rows[idx // 2])

    p = doc.add_paragraph(); set_para(p, before=7, after=4)
    add_text(p, "反馈也会有一点变化", 10.5, True, PURPLE_DARK)
    add_bullet(doc, "打卡成功", "“收到，我知道你在啦。” / “好，今天这一格也点亮了。”", GREEN)
    add_bullet(doc, "开始计时", "“那就开始吧，我陪你。” / “专心去吧，结束时再叫我。”", PURPLE)
    add_bullet(doc, "结束一段学习", "“这一段收好啦。” / “辛苦了，先喘口气也没关系。”", PEACH)
    add_bullet(doc, "错过打卡", "“这次没等到你，下个时间点见。” / “这一格先空着，我们继续往后走。”", ROSE)
    add_note_box(doc, "她不会阴阳怪气", "错过只如实记录，不连续弹窗、不制造愧疚。亲近感来自陪伴和变化，不来自压力。", CREAM, PURPLE)

    # Page 4: action language
    add_page(doc, "04")
    add_label(doc, "ACTIONS", "小鹿的动作就是她的语言")
    p = doc.add_paragraph(); set_para(p, after=5, line=1.25)
    add_text(p, "动作不只是装饰。看到她的样子，就能大概知道现在发生了什么。", 8.3)
    action_table = doc.add_table(rows=3, cols=3)
    action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    action_items = [
        (sprites["look"], "看向你", "平常待机 / 暂停", CREAM),
        (sprites["focus"], "认真进行中", "学习计时正在累计", "E8F1E5"),
        (sprites["waiting"], "等你回应", "打卡窗口已打开", LILAC),
        (sprites["wave"], "挥挥手", "打卡成功 / 开始计时", PEACH),
        (sprites["review"], "回顾一下", "结束一段 / 填今日结算", CREAM),
        (sprites["failed"], "有点失落", "错过打卡 / 约定未完成", "FCE7E8"),
        (sprites["jump"], "开心跳起", "两个人都完成约定", "E8F1E5"),
        (sprites["run_left"], "向左小跑", "拖动时朝左移动", LILAC),
        (sprites["run_right"], "向右小跑", "拖动时朝右移动", PEACH),
    ]
    for cell, item in zip((c for row in action_table.rows for c in row.cells), action_items):
        add_action_card(cell, *item)
    p = doc.add_paragraph(); set_para(p, before=5)
    add_text(p, "21 点结算后，她会回到普通待机。之后每隔几分钟短暂重现一次当日结算动作，再安静下来，不会连续跳动或卡在奇怪的一帧。", 7.8, True, PURPLE_DARK)

    # Page 5: settlement and the shared bookmark
    add_page(doc, "05")
    add_label(doc, "DAILY LOG", "21 点，一起把今天收进日记")
    p = doc.add_paragraph(); set_para(p, after=5, line=1.25)
    add_text(p, "结束打卡后，会打开一张很短的今日表单。系统自动填入学习时长和五次打卡；只需要补几项真正有意义的内容。", 8.2)

    fields = doc.add_table(rows=2, cols=2)
    fields.alignment = WD_TABLE_ALIGNMENT.CENTER
    field_items = [
        ("自动记录", "今日学习时长 · 五次打卡结果", LILAC),
        ("我来填写", "今日做题数量 · 一句成果", CREAM),
        ("我的约定", "我今天是否完成：对勾 / 叉", "E8F1E5"),
        ("你的约定", "现实里的你是否完成：对勾 / 叉", PEACH),
    ]
    for cell, (title, body, fill) in zip((c for row in fields.rows for c in row.cells), field_items):
        clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 10)
        p = cell.paragraphs[0]; set_para(p, after=2)
        add_text(p, title, 8.5, True, PURPLE_DARK)
        p = cell.add_paragraph(); set_para(p, line=1.2)
        add_text(p, body, 7.4, False, INK)

    p = doc.add_paragraph(); set_para(p, before=7, after=3)
    add_text(p, "结算只决定双人书签", 10.7, True, PURPLE_DARK)
    rewards = doc.add_table(rows=2, cols=2)
    rewards.alignment = WD_TABLE_ALIGNMENT.CENTER
    reward_items = [
        ("我们都完成", "获得一枚双人书签", GREEN),
        ("一人或两人未完成", "当天留档，不生成失败书签", "FCE7E8"),
    ]
    for row, (condition, result, fill) in zip(rewards.rows, reward_items):
        no_split(row)
        for idx, text in enumerate((condition, result)):
            cell = row.cells[idx]; clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 8)
            p = cell.paragraphs[0]; set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(p, text, 7.8, idx == 1, PURPLE_DARK if idx == 1 else INK)

    add_note_box(doc, "由我们亲自判断", "软件不会用学习时长替我们裁定结果。现实里是否守住约定，由我在结算时分别选择；两项都完成，才把共同的一天收成双人书签。", CREAM, PURPLE)

    # Page 6: daily bounties
    add_page(doc, "06")
    add_label(doc, "BOUNTY", "把想坚持的事写进两枚书签")
    p = doc.add_paragraph(); set_para(p, after=5, line=1.3)
    add_text(p, "任务页最上方的“今日悬赏”有两份固定目标。它们每天重新等待完成，但目标文字会保留，可以长期守住同一件事，也可以随时改写。", 8.3)

    bounty = doc.add_table(rows=2, cols=1)
    bounty.alignment = WD_TABLE_ALIGNMENT.CENTER
    bounty_items = [
        (BOOKMARKS / "bookmark-friend-bounty.png", "今日挑战", "为她赢下一枚书签：选择一件有挑战、但值得长期努力的事。", PEACH, 9.1),
        (BOOKMARKS / "bookmark-self-bounty.png", "今日坚持", "为自己赢下一枚书签：选择一件总想坚持、却容易放下的事。", LILAC, 8.3),
    ]
    for row, (image, title, body, fill, width) in zip(bounty.rows, bounty_items):
        cell = row.cells[0]; clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 10)
        p = cell.paragraphs[0]; set_para(p, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, title + "  ", 9.2, True, PURPLE_DARK)
        add_text(p, body, 7.4, False, INK)
        p = cell.add_paragraph(); set_para(p, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(str(image), width=Cm(width))
        no_split(row)

    p = doc.add_paragraph(); set_para(p, before=6, after=3)
    add_text(p, "直接在书签上完成它", 10.7, True, PURPLE_DARK)
    add_bullet(doc, "单击书签", "直接输入或修改目标；把文字清空，就会恢复成“尚未填写”。", PURPLE)
    add_bullet(doc, "双击书签", "完成后书签会晃一晃并收起，右上角出现 +1，收藏数量随之增加。", GREEN)
    add_bullet(doc, "想要撤回", "完成寄语右侧有一个回转箭头，点一下即可恢复今天的悬赏。", PEACH)
    add_note_box(doc, "没有失败书签", "没完成不会扣除任何东西，也不会留下刺眼的失败标记。悬赏只负责奖励主动赢下的一天。", CREAM, PURPLE)

    # Page 7: tasks and diary navigation
    add_page(doc, "07")
    add_label(doc, "TASKS", "今日清单，够用就好")
    p = doc.add_paragraph(); set_para(p, after=5, line=1.3)
    add_text(p, "悬赏下面是普通任务。它适合收下今天要做的小事，不承担复杂的项目管理，也不要求把一天塞得很满。", 8.3)

    task_table = doc.add_table(rows=3, cols=2)
    task_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    task_items = [
        ("添加与编辑", "新任务直接加入清单；文字就在任务框内修改，不需要额外的“改”按钮。", LILAC),
        ("完成与撤销", "点击左侧像素方框切换状态；误点后可以再点一次撤回。", "E8F1E5"),
        ("固定每日任务", "点亮“日”标记后，它会在第二天自动恢复为待完成。", PEACH),
        ("普通任务", "完成后仍会保留到当天结束，第二天不会无限堆积已完成项目。", CREAM),
        ("晚间提醒", "21 点后若仍有任务未完成，小鹿会用简短气泡适时提醒。", "FCE7E8"),
        ("像素翻页", "清单每页 3 项；记录每页 4 天。内容再多也不出现滚动条。", LILAC),
    ]
    for cell, (title, body, fill) in zip((c for row in task_table.rows for c in row.cells), task_items):
        clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 9)
        p = cell.paragraphs[0]; set_para(p, after=2)
        add_text(p, title, 8.5, True, PURPLE_DARK)
        p = cell.add_paragraph(); set_para(p, line=1.2)
        add_text(p, body, 7.1, False, INK)

    p = doc.add_paragraph(); set_para(p, before=7, after=3)
    add_text(p, "右键打开的四个页面", 10.7, True, PURPLE_DARK)
    pages = doc.add_table(rows=1, cols=4)
    pages.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, (title, body, fill) in zip(pages.rows[0].cells, [
        ("今日", "计时、打卡与结算", CREAM),
        ("任务", "悬赏与今日清单", PEACH),
        ("记录", "每天的学习留档", LILAC),
        ("统计", "长期累计成果", "E8F1E5"),
    ]):
        clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 9)
        p = cell.paragraphs[0]; set_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, title, 8.7, True, PURPLE_DARK)
        p = cell.add_paragraph(); set_para(p, line=1.16, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, body, 6.8, False, MUTED)

    # Page 8: bookmark collection
    add_page(doc, "08")
    add_label(doc, "COLLECTION", "三种书签，三种认真")
    p = doc.add_paragraph(); set_para(p, after=5, line=1.3)
    add_text(p, "点击日记右上角的书签图标，就能打开收藏页。每种书签只展示一枚完整图案，旁边的 × 数量记录它被赢得了多少次。", 8.3)

    gallery = doc.add_table(rows=1, cols=3)
    gallery.alignment = WD_TABLE_ALIGNMENT.CENTER
    gallery_items = [
        (BOOKMARKS / "bookmark-self.png", "为自己赢得", "完成“今日坚持”", LILAC, 6.5),
        (BOOKMARKS / "bookmark-together.png", "双人书签", "21 点两人都履约", CREAM, 7.4),
        (BOOKMARKS / "bookmark-friend.png", "为她赢得", "完成“今日挑战”", PEACH, 6.7),
    ]
    for cell, (image, title, body, fill, height) in zip(gallery.rows[0].cells, gallery_items):
        clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 10)
        p = cell.paragraphs[0]; set_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(str(image), height=Cm(height))
        p = cell.add_paragraph(); set_para(p, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, title, 8.5, True, PURPLE_DARK)
        p = cell.add_paragraph(); set_para(p, line=1.16, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, body, 6.8, False, MUTED)

    add_note_box(doc, "收藏的规则", "两枚单人书签现在都是悬赏奖励，不再代表“另一个人失败”。双人书签只记录共同完成。它们不会消费，也没有连续天数压力。", CREAM, PURPLE)

    # Page 9: statistics, local data and closing note
    add_page(doc, "09")
    add_label(doc, "OUR STORY", "日记只记录真正值得留下的事")
    stats = doc.add_table(rows=2, cols=3)
    stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    stat_items = [
        ("累计学习时长", "每次双击开启的学习段", LILAC),
        ("累计做题数量", "来自每日结算填写", CREAM),
        ("按时打卡", "五个时间点的到场记录", PEACH),
        ("双人书签", "只统计共同履约的书签", "E8F1E5"),
        ("累计完成任务", "普通任务与每日任务", LILAC),
        ("累计完成悬赏", "两份书签悬赏的总数", CREAM),
    ]
    for cell, (title, body, fill) in zip((c for row in stats.rows for c in row.cells), stat_items):
        clear_cell(cell); set_cell_shading(cell, fill); set_cell_border(cell, PURPLE_DARK, 9)
        p = cell.paragraphs[0]; set_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, title, 8.2, True, PURPLE_DARK)
        p = cell.add_paragraph(); set_para(p, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, body, 6.8, False, MUTED)

    p = doc.add_paragraph(); set_para(p, before=7, after=3)
    add_text(p, "这套系统刻意保持简单", 10.7, True, PURPLE_DARK)
    add_bullet(doc, "不记录暂停时长", "不学习的间隔不重要；只有主动双击开启的时间才累计。", GREEN)
    add_bullet(doc, "不追逐连续天数", "偶尔中断不会变成需要找补的压力，累计完成永远保留。", PEACH)
    add_bullet(doc, "数据只留在本机", "没有账号、排行榜或联网同步；所有学习记录都保存在这台电脑里。", PURPLE)

    closing = doc.add_table(rows=1, cols=1)
    closing.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = closing.cell(0, 0); clear_cell(cell); set_cell_shading(cell, PURPLE_DARK); set_cell_border(cell, PURPLE_DARK, 16)
    p = cell.paragraphs[0]; set_para(p, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "最后想告诉你", 9, True, YELLOW)
    p = cell.add_paragraph(); set_para(p, line=1.4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "电脑里的小鹿不会代替现实里的你。\n她只是把我们的约定变得看得见，也把那些认真过的日子好好留下来。", 8.7, True, WHITE)
    p = doc.add_paragraph(); set_para(p, before=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(sprites["idle"]), height=Cm(2.7))
    p = doc.add_paragraph(); set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "明天也一起吧。", 10.5, True, PURPLE_DARK)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
